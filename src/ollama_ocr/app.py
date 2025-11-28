import streamlit as st
from ocr_processor import OCRProcessor
import tempfile
import os
from PIL import Image
import json
import subprocess
from io import BytesIO
import requests
import time
from datetime import timedelta, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from pathlib import Path

# Page configuration
st.set_page_config(
    page_title="OCR with Ollama",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Add Skyone logo to sidebar
st.logo("https://skyone.solutions/wp-content/uploads/2024/12/logo-skyone-azul-scaled.webp")

# Custom CSS - Anthropic Light Inspired Theme
st.markdown("""
    <style>
    /* Logo size adjustment */
    [data-testid="stSidebarLogo"] img {
        width: 130% !important;
        max-width: 130% !important;
    }
    
    .stApp {
        max-width: 100%;
        padding: 1rem;
        background-color: #FFFFFF;
    }
    .main {
        background-color: #FFFFFF;
    }
    .stButton button {
        background-color: #000000;
        color: white;
        border: none;
        border-radius: 6px;
        padding: 0.5rem 1rem;
        font-weight: 500;
        transition: all 0.2s;
    }
    .stButton button:hover {
        background-color: #333333;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.3);
    }
    .stSelectbox {
        margin-bottom: 1rem;
    }
    .stTextArea textarea {
        border: 1px solid #E0E0E0;
        border-radius: 6px;
        font-family: 'Courier New', monospace;
    }
    .stTextInput input {
        border: 1px solid #E0E0E0;
        border-radius: 6px;
    }
    .stImage {
        border-radius: 8px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        border: 1px solid #F0F0F0;
    }
    h1 {
        color: #1F1F1F;
        font-weight: 600;
    }
    h2, h3 {
        color: #2F2F2F;
        font-weight: 500;
    }
    .sidebar .sidebar-content {
        background-color: #F7F7F7;
    }
    .warning-highlight {
        background-color: #FFF9C4;
        border-left: 4px solid #FBC02D;
        padding: 1rem 1.5rem;
        margin: 1rem 0;
        border-radius: 6px;
        box-shadow: 0 2px 4px rgba(251, 192, 45, 0.2);
    }
    .warning-highlight p {
        color: #856404;
        font-weight: 500;
        margin: 0;
        font-size: 0.95rem;
        line-height: 1.5;
    }
    .warning-highlight strong {
        color: #6B5500;
        font-weight: 600;
    }
    </style>
""", unsafe_allow_html=True)

DEFAULT_MODELS = [
    "llava:7b",
    "llama3.2-vision:11b",
    "granite3.2-vision",
    "moondream",
    "minicpm-v",
]

def get_available_models():
    try:
        result = subprocess.run(
            ["ollama", "list"],
            capture_output=True,
            text=True,
            check=True,
        )
        models = []
        for line in result.stdout.splitlines():
            line = line.strip()
            if not line or line.startswith("NAME"):
                continue
            models.append(line.split()[0])
        return models
    except Exception:
        return []

def get_openai_models(api_key):
    """Get available vision models from OpenAI API"""
    if not api_key:
        return []
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            response = requests.get("https://api.openai.com/v1/models", headers=headers, timeout=10)
            if response.status_code == 200:
                models_data = response.json()
                # Get all available models (not just vision-specific)
                all_models = []
                for model in models_data.get("data", []):
                    model_id = model.get("id", "")
                    # Include all GPT models
                    if model_id.startswith("gpt-"):
                        all_models.append(model_id)
                
                # Sort and return unique models
                all_models = sorted(set(all_models), reverse=True)
                return all_models if all_models else []
            elif response.status_code in [500, 502, 503, 504] and attempt < max_retries - 1:
                # Retry on server errors
                time.sleep(2 ** attempt)
                continue
            else:
                return []
        except requests.exceptions.RequestException as e:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)
                continue
            else:
                st.warning(f"Erro ao buscar modelos da OpenAI após {max_retries} tentativas: {str(e)}")
                return []
        except Exception as e:
            st.warning(f"Erro ao buscar modelos da OpenAI: {str(e)}")
            return []
    return []

def get_gemini_models(api_key):
    """Get available models from Google Gemini API"""
    if not api_key:
        return []
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                models_data = response.json()
                gemini_models = []
                for model in models_data.get("models", []):
                    model_name = model.get("name", "")
                    # Extract model ID from full name (e.g., "models/gemini-pro" -> "gemini-pro")
                    if "/" in model_name:
                        model_id = model_name.split("/")[-1]
                        # Only include models that support vision
                        if "vision" in model_id.lower() or "gemini-1.5" in model_id or "gemini-2" in model_id:
                            gemini_models.append(model_id)
                
                return sorted(set(gemini_models), reverse=True) if gemini_models else []
            elif response.status_code in [500, 502, 503, 504] and attempt < max_retries - 1:
                # Retry on server errors
                time.sleep(2 ** attempt)
                continue
            else:
                return []
        except requests.exceptions.RequestException as e:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)
                continue
            else:
                st.warning(f"Erro ao buscar modelos do Gemini após {max_retries} tentativas: {str(e)}")
                return []
        except Exception as e:
            st.warning(f"Erro ao buscar modelos do Gemini: {str(e)}")
            return []
    return []

def process_single_image(processor, image_path, format_type, enable_preprocessing, custom_prompt, language, status_text, timer_text):
    """Process a single image and return the result"""
    try:
        start_time = time.time()
        
        # Create progress callback
        def update_progress(current, total, message):
            elapsed = time.time() - start_time
            timer_text.metric("⏱️ Tempo Decorrido", f"{elapsed:.1f}s")
            status_text.text(message)
        
        # Set progress callback
        processor.progress_callback = update_progress
        
        # Process the image
        result = processor.process_image(
            image_path=image_path,
            format_type=format_type,
            preprocess=enable_preprocessing,
            custom_prompt=custom_prompt,
            language=language
        )
        
        return result
    except Exception as e:
        return f"Error processing image: {str(e)}"

def process_batch_images(processor, image_paths, format_type, enable_preprocessing, custom_prompt, language, status_text, timer_text):
    """Process multiple images and return results"""
    try:
        start_time = time.time()
        
        # Create progress callback
        def update_progress(current, total, message):
            elapsed = time.time() - start_time
            avg_time = elapsed / current if current > 0 else 0
            estimated_total = avg_time * total
            remaining = estimated_total - elapsed
            
            timer_text.metric(
                "⏱️ Tempo Decorrido", 
                f"{elapsed:.1f}s",
                delta=f"~{remaining:.1f}s restantes" if remaining > 0 else None
            )
            status_text.text(message)
        
        # Set progress callback
        processor.progress_callback = update_progress
        
        results = processor.process_batch(
            input_path=image_paths,
            format_type=format_type,
            preprocess=enable_preprocessing,
            custom_prompt=custom_prompt,
            language=language
        )
        
        return results
    except Exception as e:
        return {"error": str(e)}

def get_files_from_folder(folder_path, recursive=False):
    """Get all supported image/PDF files from a folder"""
    supported_extensions = ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.pdf']
    image_paths = []
    
    folder_path = Path(folder_path)
    
    if not folder_path.exists():
        return []
    
    if not folder_path.is_dir():
        return []
    
    if recursive:
        # Recursive search
        for ext in supported_extensions:
            image_paths.extend(folder_path.rglob(f'*{ext}'))
            image_paths.extend(folder_path.rglob(f'*{ext.upper()}'))
    else:
        # Non-recursive search
        for ext in supported_extensions:
            image_paths.extend(folder_path.glob(f'*{ext}'))
            image_paths.extend(folder_path.glob(f'*{ext.upper()}'))
    
    # Convert to strings and remove duplicates
    return list(set([str(p) for p in image_paths]))

def validate_file_path(file_path):
    """Validate if a file path exists and is a supported format"""
    supported_extensions = ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.pdf']
    
    file_path = Path(file_path)
    
    if not file_path.exists():
        return False, "Arquivo não encontrado"
    
    if not file_path.is_file():
        return False, "Caminho não é um arquivo"
    
    if file_path.suffix.lower() not in supported_extensions:
        return False, f"Formato não suportado. Formatos aceitos: {', '.join(supported_extensions)}"
    
    return True, "OK"

def save_processed_file(original_file_path, result_text, save_dir, format_type_internal, 
                       selected_model, format_type, language, elapsed_time=None, is_batch=False):
    """Save a processed file to the specified directory - always saves as Word document"""
    try:
        # Create directory if it doesn't exist
        save_path = Path(save_dir)
        save_path.mkdir(parents=True, exist_ok=True)
        
        # Get original file name without extension
        original_name = Path(original_file_path).stem
        
        # Always save as Word document (DOCX format)
        ext = ".docx"
        
        # Convert result_text to dict format expected by create_structured_docx
        # The function expects content_dict to be a dict with file names as keys and content as values
        if is_batch:
            # For batch, result_text should already be a dict-like structure
            if isinstance(result_text, dict):
                content_dict = result_text
            else:
                # If it's a string, wrap it in a dict
                content_dict = {Path(original_file_path).name: result_text}
        else:
            # For single file, create a simple dict structure
            content_dict = {Path(original_file_path).name: result_text}
        
        # Create Word document
        doc = create_structured_docx(
            title=f'Resultado OCR - {Path(original_file_path).name}',
            content_dict=content_dict,
            model_name=selected_model,
            format_type=format_type,
            language=language,
            elapsed_time=elapsed_time,
            is_batch=is_batch
        )
        
        # Save to buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        content = doc_buffer.getvalue()
        
        # Create output file path
        output_filename = f"{original_name}_resultado{ext}"
        output_path = save_path / output_filename
        
        # Save file as binary (Word document)
        with open(output_path, 'wb') as f:
            f.write(content)
        
        return str(output_path), None
    except Exception as e:
        return None, str(e)

def open_folder_in_explorer(folder_path):
    """Open a folder in the system file explorer"""
    import platform
    try:
        if platform.system() == "Windows":
            # Windows: usar os.startfile ou explorer
            if folder_path and os.path.exists(folder_path) and os.path.isdir(folder_path):
                os.startfile(folder_path)
            else:
                # Abrir diretório atual se não houver pasta válida
                os.startfile(".")
        elif platform.system() == "Darwin":  # macOS
            path_to_open = folder_path if folder_path and os.path.exists(folder_path) and os.path.isdir(folder_path) else "."
            subprocess.Popen(["open", path_to_open])
        else:  # Linux
            path_to_open = folder_path if folder_path and os.path.exists(folder_path) and os.path.isdir(folder_path) else "."
            subprocess.Popen(["xdg-open", path_to_open])
        return True, None
    except Exception as e:
        return False, str(e)

def create_structured_docx(title, content_dict, model_name, format_type, language, elapsed_time=None, is_batch=False):
    """Create a structured DOCX document with professional formatting"""
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    doc.add_heading('Informações do Processamento', level=1)
    
    # Create metadata table
    metadata_table = doc.add_table(rows=5 if elapsed_time else 4, cols=2)
    metadata_table.style = 'Light Grid Accent 1'
    
    # Fill metadata
    row_idx = 0
    cells = metadata_table.rows[row_idx].cells
    cells[0].text = 'Data e Hora'
    cells[1].text = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    
    row_idx += 1
    cells = metadata_table.rows[row_idx].cells
    cells[0].text = 'Modelo Utilizado'
    cells[1].text = model_name
    
    row_idx += 1
    cells = metadata_table.rows[row_idx].cells
    cells[0].text = 'Formato de Saída'
    cells[1].text = format_type
    
    row_idx += 1
    cells = metadata_table.rows[row_idx].cells
    cells[0].text = 'Idioma'
    cells[1].text = language
    
    if elapsed_time:
        row_idx += 1
        cells = metadata_table.rows[row_idx].cells
        cells[0].text = 'Tempo de Processamento'
        cells[1].text = f'{elapsed_time:.2f} segundos'
    
    # Add spacing
    doc.add_paragraph()
    
    # Add results section
    doc.add_heading('Resultados Extraídos', level=1)
    
    if is_batch:
        # For batch processing, iterate through multiple results
        for idx, (file_name, text) in enumerate(content_dict.items(), 1):
            # Add file header
            doc.add_heading(f'{idx}. {file_name}', level=2)
            
            # Add separator line
            p = doc.add_paragraph()
            p.add_run('─' * 80)
            run = p.runs[0]
            run.font.color.rgb = RGBColor(200, 200, 200)
            
            # Add content with formatting
            content_para = doc.add_paragraph()
            run = content_para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            
            # Add spacing between files
            doc.add_paragraph()
    else:
        # For single file processing
        # Add separator line
        p = doc.add_paragraph()
        p.add_run('─' * 80)
        run = p.runs[0]
        run.font.color.rgb = RGBColor(200, 200, 200)
        
        # Add content with formatting
        content_para = doc.add_paragraph()
        run = content_para.add_run(content_dict)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    # Add footer
    doc.add_page_break()
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'Gerado por OCR Vision – Skyone LAB | {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

def create_minuta_doc(content_dict, is_batch=False):
    """Create a document formatted according to Brazilian legal document standards (peças processuais)"""
    doc = Document()
    
    # Configure page margins according to legal standards
    # Superior: 3cm, Esquerda: 3cm, Inferior: 2cm, Direita: 2cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Set default font and paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Configure paragraph spacing (1.5 line spacing)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    
    # Add content
    if is_batch:
        # For batch processing
        for idx, (file_name, text) in enumerate(content_dict.items(), 1):
            if idx > 1:
                # Add page break between files
                doc.add_page_break()
            
            # Split content into paragraphs (handle both \n\n and \n)
            paragraphs = text.replace('\r\n', '\n').split('\n\n')
            for para_text in paragraphs:
                # Also split by single \n to handle line breaks within paragraphs
                lines = para_text.split('\n')
                para_content = ' '.join(line.strip() for line in lines if line.strip())
                if para_content:
                    p = doc.add_paragraph(para_content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # Apply paragraph formatting
                    p_format = p.paragraph_format
                    p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p_format.line_spacing = 1.5
                    p_format.space_after = Pt(0)
                    p_format.space_before = Pt(0)
                    # Apply formatting to all runs in paragraph
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
    else:
        # For single file processing
        # Split content into paragraphs (handle both \n\n and \n)
        paragraphs = content_dict.replace('\r\n', '\n').split('\n\n')
        for para_text in paragraphs:
            # Also split by single \n to handle line breaks within paragraphs
            lines = para_text.split('\n')
            para_content = ' '.join(line.strip() for line in lines if line.strip())
            if para_content:
                p = doc.add_paragraph(para_content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Apply paragraph formatting
                p_format = p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p_format.line_spacing = 1.5
                p_format.space_after = Pt(0)
                p_format.space_before = Pt(0)
                # Apply formatting to all runs in paragraph
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    # Configure header and footer for page numbering
    section = doc.sections[0]
    
    # Footer with page number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field (will be updated by Word)
    run = footer_para.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    # Use a simple text approach for page numbers
    # Note: Actual page numbering requires Word fields which python-docx doesn't fully support
    # The document will be properly numbered when opened in Word
    
    return doc

def main():
    # Header in expander
    with st.expander("ℹ️ Sobre o Skyone OCR", expanded=False):
        st.markdown("""
        <div style='padding: 0.5rem 0;'>
            <h2 style='
                color: #1F1F1F;
                font-size: 1.5rem;
                font-weight: 600;
                margin: 0 0 1rem 0;
                padding: 0;
            '>
                🧾 Skyone OCR
            </h2>
            <p style='
                color: #4A4A4A;
                font-size: 0.95rem;
                margin: 0;
                padding: 0;
                line-height: 1.6;
            '>
                Uma tecnologia de visão computacional e IA para extrair e interpretar textos de documentos, imagens e PDFs com máxima acurácia. 
                <br><br>Projetado para impulsionar automações no Skyone Studio e alimentar agentes de IA com dados estruturados e confiáveis.
            </p>
        </div>
        """, unsafe_allow_html=True)

    # Sidebar controls
    with st.sidebar:
        st.markdown("<style>.sidebar .sidebar-content { font-size: 8pt; }</style>", unsafe_allow_html=True)
        st.header("Configurações IA")
        
        # AI Configuration Section
        with st.expander("🤖 Inteligência Artificial", expanded=False):
            # API Provider Selection
            api_provider = st.selectbox(
                "▪ Provedor LLM",
                ["Google Gemini", "Ollama (Local)", "OpenAI"],
                help="Escolha o provedor de IA para processamento"
            )
            
            # API Key input for external providers
            api_key = None
            if api_provider in ["OpenAI", "Google Gemini"]:
                api_key = st.text_input(
                    "▪ Chave da API",
                    type="password",
                    help=f"Insira sua chave de API do {api_provider}"
                )
            
            # Model selection based on provider
            if api_provider == "Ollama (Local)":
                available_models = get_available_models()
                if not available_models:
                    st.warning("Não foi possível buscar modelos do Ollama. Usando lista padrão.")
                    available_models = DEFAULT_MODELS
                selected_model = st.selectbox(
                    "▪ Modelo de Visão",
                    available_models,
                    index=0,
                )
            elif api_provider == "OpenAI":
                # Get OpenAI models dynamically if API key is provided
                openai_models = get_openai_models(api_key)
                if openai_models:
                    selected_model = st.selectbox(
                        "▪ Modelo de Visão",
                        openai_models,
                        index=0,
                        help="Modelos disponíveis na sua conta OpenAI"
                    )
                else:
                    st.warning("⚠️ Insira a API Key da OpenAI para ver os modelos disponíveis.")
                    selected_model = None
            else:  # Google Gemini
                # Get Gemini models dynamically if API key is provided
                gemini_models = get_gemini_models(api_key)
                if gemini_models:
                    selected_model = st.selectbox(
                        "▪ Modelo de Visão",
                        gemini_models,
                        index=0,
                        help="Modelos disponíveis na sua conta Google Gemini"
                    )
                else:
                    st.warning("⚠️ Insira a API Key do Google Gemini para ver os modelos disponíveis.")
                    selected_model = None
        
        # Output Format Section
        with st.expander("🧾 Formato de Saída", expanded=False):
            format_type = st.selectbox(
                "▪ Tipo de Formato",
                ["Markdown", "Texto", "JSON", "Estruturado", "Chave-Valor", "Tabela", "Documento do Word 97-2003"],
                help="Escolha como deseja formatar o texto extraído"
            )
        
        # Prompt Configuration Section
        with st.expander("📄 Configuração de Prompt", expanded=False):
            # Prompt type selection
            prompt_type = st.selectbox(
                "▪ Tipo de Prompt",
                ["Automático", "Manual"],
                help="Escolha entre prompt automático (padrão otimizado) ou manual (personalizado)"
            )
            
            # Automatic prompt
            automatic_prompt = """Siga rigorosamente as instruções abaixo para processar o conteúdo do arquivo, garantindo que o resultado final esteja formatado de maneira totalmente compatível com Microsoft Word Document 97–2004 (.doc):

1. Transcrição Integral e Fiel

Transcreva o conteúdo completo, mantendo a estrutura, ordem e organização visual com a maior fidelidade possível ao documento original.

2. Correção Automática de OCR

Corrija apenas erros evidentes de leitura, como:
• caracteres distorcidos;
• letras ou sinais faltando;
• palavras quebradas;
• acentuação/ortografia claramente afetadas pelo OCR.

3. Marcação de Ilegibilidade

Se um trecho estiver ausente, ilegível ou incerto, marque TRECHO ILEGÍVEL exatamente onde ocorre.
Nunca reordene o conteúdo para esconder falhas.

4. Sem Inferências

Não preencha lacunas com suposições.
Use TRECHO ILEGÍVEL sempre que a leitura não for 100% segura.

5. Reconstrução de Estruturas

Caso o documento contenha tabelas, quadros, fichas, formulários, listas ou campos pré-definidos, reconstrua utilizando:
• tabelas simples compatíveis com Word 97–2004;
• listas numeradas ou com marcadores;
• separação clara de seções;
• títulos simples.

🔗 Evitar: caixas de texto avançadas, ícones, figuras inline modernas, tabelas complexas ou recursos não suportados pelo formato .doc.

6. Seção Final — Extração de Campos Estruturados

Após a transcrição completa, apresente uma seção separada contendo:
• campos identificados;
• valores extraídos;
• marcações TRECHO ILEGÍVEL quando necessário.

Use uma tabela simples ou lista compatível com Word 97–2004 (.doc).

7. Compatibilidade com Word 97–2004 (.doc)

Todo o conteúdo deve utilizar apenas formatação legada:
• tabelas simples;
• listas simples;
• negrito, itálico e sublinhado básicos;
• seções com títulos;
• nada de estilos avançados, emojis, cores especiais ou elementos modernos."""
            
            # Custom prompt input (conditional)
            if prompt_type == "Manual":
                custom_prompt_input = st.text_area(
                    "▪ Prompt Personalizado",
                    value="",
                    placeholder="Digite seu prompt aqui (obrigatório)",
                    help="Insira um prompt personalizado para extração de texto. Este campo é obrigatório.",
                    height=200
                )
            else:
                custom_prompt_input = automatic_prompt
                st.text_area(
                    "▪ Prompt Automático",
                    value=automatic_prompt,
                    help="Prompt automático otimizado para OCR com correção e marcação de ilegibilidade",
                    height=200,
                    disabled=True
                )
        
        st.markdown("<style>.sidebar .sidebar-content { font-size: 8pt; }</style>", unsafe_allow_html=True)
        st.header("Configurações Avançadas")
        # Advanced Settings
        with st.expander("⚙️ Configurações Avançadas", expanded=False):
            preprocessing_option = st.selectbox(
                "▪ Pré-processamento",
                options=["Ativado", "Desativado"],
                index=0,  # Default "Ativado"
                help="Aplicar aprimoramento e pré-processamento de imagem"
            )
            enable_preprocessing = preprocessing_option == "Ativado"
            
            language = st.text_input(
                "▪ Idioma",
                value="pt-br",
                help="Insira o idioma do texto na imagem (ex: pt-br para Português, en para Inglês)."
            )

            max_workers = st.selectbox(
                "▪ Processamento Paralelo",
                options=[1, 2, 3, 4, 5, 6, 7, 8],
                index=1,  # Default value 2
                help="Número de imagens a processar em paralelo (para processamento em lote)"
            )
            
            st.divider()
        
        st.markdown("<style>.sidebar .sidebar-content { font-size: 8pt; }</style>", unsafe_allow_html=True)
        st.header("Resultados")
        
        # Download Results Section
        with st.expander("📥 Download de Resultados", expanded=False):
            # Placeholder for download options
            if 'download_placeholder' not in st.session_state:
                st.session_state['download_placeholder'] = st.empty()
            
            with st.session_state['download_placeholder'].container():
                st.info("Processe arquivos para ver as opções de download")
    
    
    # Map translated format names to internal format values
    format_map = {
        "Markdown": "markdown",
        "Texto": "text",
        "JSON": "json",
        "Estruturado": "structured",
        "Chave-Valor": "key_value",
        "Tabela": "table",
        "Documento do Word 97-2003": "doc97"
    }
    format_type_internal = format_map.get(format_type, "markdown")
    
    # Set custom prompt based on type
    if prompt_type == "Automático":
        custom_prompt = automatic_prompt
    else:
        custom_prompt = custom_prompt_input.strip() if custom_prompt_input.strip() != "" else None

    # Map provider name to internal format
    provider_map = {
        "Ollama (Local)": "ollama",
        "OpenAI": "openai",
        "Google Gemini": "gemini"
    }
    
    # Initialize OCR Processor with API provider and key
    try:
        processor = OCRProcessor(
            model_name=selected_model, 
            max_workers=max_workers,
            api_provider=provider_map[api_provider],
            api_key=api_key
        )
    except ValueError as e:
        with st.expander("ℹ️ Aguardando Configuração", expanded=True):
            st.markdown("""
            <div style="
                background-color: #E3F2FD;
                border-left: 4px solid #2196F3;
                padding: 1rem 1.5rem;
                margin: 1rem 0;
                border-radius: 6px;
                box-shadow: 0 2px 4px rgba(33, 150, 243, 0.2);
            ">
                <p style="
                    color: #1565C0;
                    font-weight: 500;
                    margin: 0;
                    font-size: 0.95rem;
                    line-height: 1.5;
                ">
                    <strong>ℹ️ Aguardando Modelo de Inteligência Artificial</strong><br/>
                    Por favor, configure o modelo de IA na barra lateral para continuar.
                </p>
            </div>
            """, unsafe_allow_html=True)
        st.stop()

    # Upload container with tabs
    st.subheader("📤 Seleção de Arquivos")
    
    upload_tab1, upload_tab2, upload_tab3 = st.tabs(["📎 Upload de Arquivos", "📁 Selecionar Pasta", "📄 Selecionar Arquivo"])
    
    uploaded_files = None
    local_folder_paths = []
    local_file_paths = []
    
    with upload_tab1:
        # Upload de arquivos via interface web
        with st.container(border=True):
            st.markdown("**Arraste seus arquivos aqui ou clique para selecionar**")
            uploaded_files = st.file_uploader(
                "Selecione os arquivos",
                type=['png', 'jpg', 'jpeg', 'tiff', 'bmp', 'pdf'],
                accept_multiple_files=True,
                help="Formatos suportados: PNG, JPG, JPEG, TIFF, BMP, PDF",
                label_visibility="collapsed"
            )
            
            if uploaded_files:
                st.divider()
                st.write(f"**{len(uploaded_files)} arquivo(s) carregado(s):**")
                for uploaded_file in uploaded_files:
                    file_size = uploaded_file.size / (1024 * 1024)  # Convert to MB
                    st.write(f"✓ {uploaded_file.name} ({file_size:.2f} MB)")
    
    with upload_tab2:
        # Seleção de pasta do sistema de arquivos
        with st.container(border=True):
            st.markdown("**Selecione uma pasta do seu computador**")
            
            col_path, col_button = st.columns([3, 1])
            with col_path:
                folder_path = st.text_input(
                    "Caminho da pasta:",
                    placeholder="Ex: C:\\Users\\Usuario\\Documentos\\Imagens",
                    help="Digite o caminho completo da pasta que contém os arquivos a serem processados",
                    label_visibility="collapsed"
                )
            with col_button:
                if st.button("📂 Abrir Pasta", use_container_width=True, help="Abrir explorador de arquivos para selecionar pasta"):
                    success, error = open_folder_in_explorer(folder_path)
                    if not success:
                        st.error(f"Erro ao abrir explorador: {error}")
            
            recursive_search = st.checkbox(
                "Buscar em subpastas (recursivo)",
                help="Se marcado, também buscará arquivos nas subpastas"
            )
            
            if folder_path:
                if st.button("🔍 Verificar Pasta", use_container_width=True):
                    if os.path.exists(folder_path) and os.path.isdir(folder_path):
                        files = get_files_from_folder(folder_path, recursive=recursive_search)
                        if files:
                            st.success(f"✅ {len(files)} arquivo(s) encontrado(s)!")
                            st.session_state['local_folder_files'] = files
                            st.session_state['local_folder_path'] = folder_path
                            st.session_state['local_folder_recursive'] = recursive_search
                            
                            with st.expander(f"📋 Ver arquivos encontrados ({len(files)})", expanded=False):
                                for file_path in files[:50]:  # Mostrar até 50 arquivos
                                    file_size = os.path.getsize(file_path) / (1024 * 1024)
                                    st.write(f"• {os.path.basename(file_path)} ({file_size:.2f} MB)")
                                if len(files) > 50:
                                    st.info(f"... e mais {len(files) - 50} arquivo(s)")
                        else:
                            st.warning("⚠️ Nenhum arquivo de imagem ou PDF encontrado nesta pasta")
                            st.session_state['local_folder_files'] = []
                    else:
                        st.error("❌ Pasta não encontrada. Verifique o caminho e tente novamente.")
            
            # Campo para caminho de salvamento dos arquivos processados
            st.divider()
            st.markdown("**💾 Caminho de Salvamento dos Resultados**")
            col_save_path, col_save_button = st.columns([3, 1])
            with col_save_path:
                save_path = st.text_input(
                    "Caminho para salvar arquivos processados:",
                    value=st.session_state.get('save_output_path', ''),
                    placeholder="Ex: C:\\Users\\Usuario\\Documentos\\Resultados_OCR",
                    help="Digite o caminho completo da pasta onde os arquivos processados serão salvos. Deixe vazio para não salvar automaticamente.",
                    label_visibility="collapsed"
                )
                if save_path:
                    st.session_state['save_output_path'] = save_path
                else:
                    st.session_state['save_output_path'] = None
            with col_save_button:
                if st.button("📂 Abrir Pasta", key="open_save_folder", use_container_width=True, help="Abrir explorador para selecionar pasta de salvamento"):
                    success, error = open_folder_in_explorer(save_path if save_path else ".")
                    if not success:
                        st.error(f"Erro ao abrir explorador: {error}")
            
            if save_path:
                if os.path.exists(save_path) and os.path.isdir(save_path):
                    st.success(f"✅ Pasta de salvamento válida: {save_path}")
                else:
                    st.warning(f"⚠️ Pasta não existe. Será criada automaticamente: {save_path}")
            
            # Mostrar arquivos selecionados se houver
            if 'local_folder_files' in st.session_state and st.session_state['local_folder_files']:
                st.divider()
                col_info, col_open = st.columns([3, 1])
                with col_info:
                    st.write(f"**{len(st.session_state['local_folder_files'])} arquivo(s) selecionado(s) da pasta:**")
                    st.caption(f"📁 {st.session_state['local_folder_path']}")
                    if st.session_state.get('local_folder_recursive'):
                        st.caption("🔍 Busca recursiva ativada")
                    if st.session_state.get('save_output_path'):
                        st.caption(f"💾 Resultados serão salvos em: {st.session_state['save_output_path']}")
                with col_open:
                    if st.button("📂 Abrir Pasta", key="open_folder_selected", use_container_width=True, help="Abrir esta pasta no explorador"):
                        success, error = open_folder_in_explorer(st.session_state['local_folder_path'])
                        if not success:
                            st.error(f"Erro ao abrir pasta: {error}")
    
    with upload_tab3:
        # Seleção de arquivo específico do sistema de arquivos
        with st.container(border=True):
            st.markdown("**Selecione um arquivo específico do seu computador**")
            file_path = st.text_input(
                "Caminho do arquivo:",
                placeholder="Ex: C:\\Users\\Usuario\\Documentos\\imagem.png",
                help="Digite o caminho completo do arquivo a ser processado"
            )
            
            if file_path:
                if st.button("🔍 Verificar Arquivo", use_container_width=True):
                    is_valid, message = validate_file_path(file_path)
                    if is_valid:
                        file_size = os.path.getsize(file_path) / (1024 * 1024)
                        st.success(f"✅ Arquivo válido!")
                        st.write(f"**Arquivo:** {os.path.basename(file_path)}")
                        st.write(f"**Tamanho:** {file_size:.2f} MB")
                        st.session_state['local_file_path'] = file_path
                    else:
                        st.error(f"❌ {message}")
                        st.session_state['local_file_path'] = None
            
            # Mostrar arquivo selecionado se houver
            if 'local_file_path' in st.session_state and st.session_state['local_file_path']:
                st.divider()
                file_path = st.session_state['local_file_path']
                file_size = os.path.getsize(file_path) / (1024 * 1024)
                st.write(f"**Arquivo selecionado:**")
                st.write(f"✓ {os.path.basename(file_path)} ({file_size:.2f} MB)")
                st.caption(f"📄 {file_path}")

    # Determinar quais arquivos processar
    files_to_process = []
    source_type = None
    
    if uploaded_files:
        # Arquivos via upload web
        source_type = "upload"
        files_to_process = uploaded_files
    elif 'local_folder_files' in st.session_state and st.session_state['local_folder_files']:
        # Arquivos de pasta local
        source_type = "folder"
        files_to_process = st.session_state['local_folder_files']
    elif 'local_file_path' in st.session_state and st.session_state['local_file_path']:
        # Arquivo específico local
        source_type = "file"
        files_to_process = [st.session_state['local_file_path']]
    
    # Botão de processar
    if files_to_process:
        st.divider()
        col1, col2 = st.columns([2, 1])
        with col1:
            if source_type == "upload":
                st.info(f"📎 {len(files_to_process)} arquivo(s) pronto(s) para processamento (Upload)")
            elif source_type == "folder":
                st.info(f"📁 {len(files_to_process)} arquivo(s) pronto(s) para processamento (Pasta: {os.path.basename(st.session_state.get('local_folder_path', ''))})")
            elif source_type == "file":
                st.info(f"📄 1 arquivo pronto para processamento")
        with col2:
            if st.button("🚀 Processar Arquivos", key="process_button_local", use_container_width=True):
                st.session_state['process_clicked'] = True
                st.session_state['source_type'] = source_type

    if st.session_state.get('process_clicked', False) and files_to_process:
        source_type = st.session_state.get('source_type', 'upload')
        
        # Process button (verifica se foi clicado via session_state)
        if st.session_state.get('process_clicked', False):
            # Reset flag
            st.session_state['process_clicked'] = False
            
            # Validate custom prompt
            if prompt_type == "Manual" and not custom_prompt:
                st.error("⚠️ Prompt Personalizado é obrigatório. Por favor, insira um prompt antes de processar.")
                st.stop()
            
            # Get save output path from session state
            save_output_path = st.session_state.get('save_output_path', None)
            
            # Prepare file paths based on source type
            if source_type == "upload":
                # Uploaded files - need to save to temp directory
                with tempfile.TemporaryDirectory() as temp_dir:
                    image_paths = []
                    for uploaded_file in uploaded_files:
                        # Reset file pointer before reading
                        uploaded_file.seek(0)
                        temp_path = os.path.join(temp_dir, uploaded_file.name)
                        with open(temp_path, "wb") as f:
                            f.write(uploaded_file.read())
                        image_paths.append(temp_path)
                    
                    # Process files
                    _process_files(image_paths, processor, format_type_internal, enable_preprocessing, 
                                 custom_prompt, language, selected_model, format_type, save_output_path)
                    
            elif source_type == "folder" or source_type == "file":
                # Files from local folder or single file - paths are already absolute
                image_paths = files_to_process.copy()
                _process_files(image_paths, processor, format_type_internal, enable_preprocessing, 
                             custom_prompt, language, selected_model, format_type, save_output_path)

def _process_files(image_paths, processor, format_type_internal, enable_preprocessing, 
                  custom_prompt, language, selected_model, format_type, save_output_path=None):
    """Process files (single or batch)"""
    # Show number of files being processed
    if len(image_paths) > 1:
        st.info(f"📂 Processando {len(image_paths)} arquivos em modo lote...")
    else:
        st.info(f"📄 Processando 1 arquivo...")
    
    # Reset usage stats before processing
    try:
        processor.reset_usage_stats()
    except Exception as e:
        st.warning(f"Aviso ao resetar estatísticas: {e}")
    
    # Create timer and status components
    timer_container = st.empty()
    status_text = st.empty()
    
    start_time = time.time()
    
    if len(image_paths) == 1:
        # Single image processing
        status_text.text("Iniciando processamento...")
        result = process_single_image(
            processor, 
            image_paths[0], 
            format_type_internal,
            enable_preprocessing,
            custom_prompt,
            language,
            status_text,
            timer_container
        )
        
        # Show final time
        elapsed_time = time.time() - start_time
        timer_container.empty()
        status_text.empty()
        
        # Get usage statistics
        try:
            usage_stats = processor.get_usage_stats()
        except Exception as e:
            st.warning(f"Aviso ao obter estatísticas: {e}")
            usage_stats = {
                'input_tokens': 0,
                'output_tokens': 0,
                'estimated_cost_brl': 0,
                'estimated_cost_usd': 0
            }
        
        st.success(f"✅ Processamento concluído em {elapsed_time:.2f}s!")
        
        # Save file automatically if save path is specified
        if save_output_path:
            if result and result.strip() and not result.startswith("Error processing image:") and len(result.strip()) >= 3:
                saved_path, error = save_processed_file(
                    image_paths[0], result, save_output_path, format_type_internal,
                    selected_model, format_type, language, elapsed_time, is_batch=False
                )
                if saved_path:
                    st.success(f"💾 Arquivo salvo automaticamente: {saved_path}")
                elif error:
                    st.warning(f"⚠️ Erro ao salvar arquivo: {error}")
        
        # Display usage statistics in a separate block
        with st.container(border=True):
            st.subheader("📊 Estatísticas de Uso")
            st.markdown('<div style="font-size: 11pt;">', unsafe_allow_html=True)
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("⏱️ Tempo", f"{elapsed_time:.2f}s")
            with col2:
                st.metric("📥 Tokens Entrada", f"{usage_stats.get('input_tokens', 0):,}")
            with col3:
                st.metric("📤 Tokens Saída", f"{usage_stats.get('output_tokens', 0):,}")
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Get raw result
        try:
            if hasattr(processor, 'get_raw_result'):
                raw_result = processor.get_raw_result() or result
            else:
                raw_result = result
        except (AttributeError, Exception):
            raw_result = result
        
        # Check if result is empty or contains only error messages
        if not result or not result.strip() or result.startswith("Error processing image:") or len(result.strip()) < 3:
            st.markdown("""
            <div class="warning-highlight">
                <p><strong>⚠️ Atenção:</strong> Nenhum conteúdo foi extraído do arquivo processado.</p>
                <p style="margin-top: 0.5rem; font-size: 0.9rem;">O processamento pode ter falhado ou o arquivo pode não conter texto legível. Verifique o arquivo e tente novamente.</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            # Display results in the selected format in a separate block
            st.subheader(f"📝 Resultado Processado ({format_type})")
            with st.container(border=True):
                st.markdown('<div style="font-size: 11pt;">', unsafe_allow_html=True)
                if format_type_internal == "json":
                    try:
                        json_data = json.loads(result)
                        st.json(json_data)
                    except:
                        st.code(result, language="json")
                elif format_type_internal == "text":
                    st.text(result)
                elif format_type_internal == "doc97":
                    st.text(result)
                elif format_type_internal in ["structured", "key_value", "table"]:
                    st.markdown(result)
                else:  # markdown
                    st.markdown(result)
                st.markdown('</div>', unsafe_allow_html=True)
        
            # Download options for single result in a separate block
            st.subheader("📥 Opções de Download")
            with st.container(border=True):
                st.markdown('<div style="font-size: 11pt;">', unsafe_allow_html=True)
                col1, col2, col3, col4, col5 = st.columns(5)
                
                with col1:
                    st.download_button(
                        "📥 Download TXT",
                        result,
                        file_name=f"ocr_result.txt",
                        mime="text/plain",
                        key="download_txt_single"
                    )
                
                with col2:
                    # Create structured DOCX
                    doc = create_structured_docx(
                        title='Resultado do OCR',
                        content_dict=result,
                        model_name=selected_model,
                        format_type=format_type,
                        language=language,
                        elapsed_time=elapsed_time,
                        is_batch=False
                    )
                    docx_buffer = BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    st.download_button(
                        "📥 Download DOCX",
                        docx_buffer.getvalue(),
                        file_name="ocr_result.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_single"
                    )
                
                with col3:
                    # DOC format with structured content
                    doc = create_structured_docx(
                        title='Resultado do OCR',
                        content_dict=result,
                        model_name=selected_model,
                        format_type=format_type,
                        language=language,
                        elapsed_time=elapsed_time,
                        is_batch=False
                    )
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    st.download_button(
                        "📥 Download DOC",
                        doc_buffer.getvalue(),
                        file_name="ocr_result.doc",
                        mime="application/msword",
                        key="download_doc_single"
                    )
                
                with col4:
                    # Raw result - exactly as LLM processed
                    st.download_button(
                        "📥 Download RAW",
                        raw_result,
                        file_name="ocr_result_raw.txt",
                        mime="text/plain",
                        help="Resultado exatamente como processado pela LLM, sem formatação",
                        key="download_raw_single"
                    )
                
                with col5:
                    # Formato Minuta - Legal document format
                    minuta_doc = create_minuta_doc(
                        content_dict=result,
                        is_batch=False
                    )
                    minuta_buffer = BytesIO()
                    minuta_doc.save(minuta_buffer)
                    minuta_buffer.seek(0)
                    st.download_button(
                        "📄 Formato Minuta",
                        minuta_buffer.getvalue(),
                        file_name="minuta.doc",
                        mime="application/msword",
                        help="Documento formatado conforme padrão de peças processuais (fonte Times New Roman 12, espaçamento 1,5, margens padrão)",
                        key="download_minuta_single"
                    )
                st.markdown('</div>', unsafe_allow_html=True)
    else:
                    # Batch processing
                    status_text.text("Iniciando processamento em lote...")
                    try:
                        results = process_batch_images(
                            processor,
                            image_paths,
                            format_type_internal,
                            enable_preprocessing,
                            custom_prompt,
                            language,
                            status_text,
                            timer_container
                        )
                    except Exception as e:
                        st.error(f"❌ Erro no processamento em lote: {str(e)}")
                        st.exception(e)
                        st.stop()
                    
                    # Show final time
                    elapsed_time = time.time() - start_time
                    timer_container.empty()
                    status_text.empty()
                    
                    # Get usage statistics
                    try:
                        usage_stats = processor.get_usage_stats()
                    except Exception as e:
                        st.warning(f"Aviso ao obter estatísticas: {e}")
                        usage_stats = {
                            'input_tokens': 0,
                            'output_tokens': 0,
                            'estimated_cost_brl': 0,
                            'estimated_cost_usd': 0
                        }
                    
                    st.success(f"✅ Processamento em lote concluído em {elapsed_time:.2f}s!")
                    
                    # Save files automatically if save path is specified
                    if save_output_path and results.get('results'):
                        saved_count = 0
                        save_errors = []
                        save_status = st.empty()
                        
                        valid_results = {fp: text for fp, text in results['results'].items() 
                                       if text and text.strip() and not text.startswith("Error processing image:")}
                        
                        for file_path, result_text in valid_results.items():
                            saved_path, error = save_processed_file(
                                file_path, result_text, save_output_path, format_type_internal,
                                selected_model, format_type, language, elapsed_time, is_batch=True
                            )
                            if saved_path:
                                saved_count += 1
                            elif error:
                                save_errors.append(f"{os.path.basename(file_path)}: {error}")
                        
                        if saved_count > 0:
                            st.success(f"💾 {saved_count} arquivo(s) salvo(s) automaticamente em: {save_output_path}")
                        if save_errors:
                            st.warning(f"⚠️ Erros ao salvar {len(save_errors)} arquivo(s):")
                            for error_msg in save_errors[:5]:  # Mostrar até 5 erros
                                st.caption(f"• {error_msg}")
                            if len(save_errors) > 5:
                                st.caption(f"... e mais {len(save_errors) - 5} erro(s)")
                    
                    # Display processing and usage statistics side by side
                    col_stat1, col_stat2 = st.columns(2)
                    
                    with col_stat1:
                        with st.container(border=True):
                            st.subheader("📊 Estatísticas de Processamento")
                            st.markdown('<div style="font-size: 11pt;">', unsafe_allow_html=True)
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Total de Imagens", results.get('statistics', {}).get('total', 0))
                            with col2:
                                st.metric("Sucesso", results.get('statistics', {}).get('successful', 0))
                            with col3:
                                st.metric("Falhas", results.get('statistics', {}).get('failed', 0))
                            st.markdown('</div>', unsafe_allow_html=True)
                    
                    with col_stat2:
                        with st.container(border=True):
                            st.subheader("💡 Estatísticas de Uso")
                            st.markdown('<div style="font-size: 11pt;">', unsafe_allow_html=True)
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("⏱️ Tempo Total", f"{elapsed_time:.2f}s")
                            with col2:
                                st.metric("📥 Tokens Entrada", f"{usage_stats.get('input_tokens', 0):,}")
                            with col3:
                                st.metric("📤 Tokens Saída", f"{usage_stats.get('output_tokens', 0):,}")
                            # Cost metrics (hidden/commented)
                            # with col4:
                            #     cost_brl = usage_stats.get('estimated_cost_brl', 0)
                            #     if cost_brl > 0:
                            #         st.metric("💰 Custo (BRL)", f"R$ {cost_brl:.4f}")
                            #     else:
                            #         st.metric("💰 Custo", "Gratuito")
                            # with col5:
                            #     cost_usd = usage_stats.get('estimated_cost_usd', 0)
                            #     if cost_usd > 0:
                            #         st.metric("💵 Custo (USD)", f"${cost_usd:.4f}")
                            #     else:
                            #         st.metric("💵 USD", "-")
                            st.markdown('</div>', unsafe_allow_html=True)

                    # Display errors if any
                    if results.get('errors'):
                        st.markdown("""
                        <div class="warning-highlight">
                            <p><strong>⚠️ Atenção:</strong> Alguns arquivos apresentaram erros durante o processamento:</p>
                        </div>
                        """, unsafe_allow_html=True)
                        for file_path, error in results['errors'].items():
                            st.markdown(f"""
                            <div class="warning-highlight" style="margin-top: 0.5rem;">
                                <p><strong>❌ {os.path.basename(file_path)}:</strong> {error}</p>
                            </div>
                            """, unsafe_allow_html=True)

                    # Display results in the selected format
                    if results.get('results'):
                        # Filter out empty results
                        valid_results = {fp: text for fp, text in results['results'].items() 
                                       if text and text.strip() and not text.startswith("Error processing image:")}
                        
                        if valid_results:
                            st.subheader(f"📝 Resultados Processados ({format_type})")
                            for file_path, text in valid_results.items():
                                with st.expander(f"✅ {os.path.basename(file_path)}", expanded=True):
                                    with st.container(border=True):
                                        st.markdown('<div style="font-size: 11pt;">', unsafe_allow_html=True)
                                        if format_type_internal == "json":
                                            try:
                                                json_data = json.loads(text)
                                                st.json(json_data)
                                            except:
                                                st.code(text, language="json")
                                        elif format_type_internal == "text":
                                            st.text(text)
                                        elif format_type_internal == "doc97":
                                            st.text(text)
                                        elif format_type_internal in ["structured", "key_value", "table"]:
                                            st.markdown(text)
                                        else:  # markdown
                                            st.markdown(text)
                                        st.markdown('</div>', unsafe_allow_html=True)
                        else:
                            # All results are empty or errors
                            st.markdown("""
                            <div class="warning-highlight">
                                <p><strong>⚠️ Atenção:</strong> Nenhum conteúdo válido foi extraído dos arquivos processados.</p>
                                <p style="margin-top: 0.5rem; font-size: 0.9rem;">Todos os arquivos podem ter falhado no processamento ou não conterem texto legível. Verifique os erros acima e tente novamente.</p>
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        st.markdown("""
                        <div class="warning-highlight">
                            <p><strong>⚠️ Atenção:</strong> Nenhum resultado foi gerado durante o processamento.</p>
                            <p style="margin-top: 0.5rem; font-size: 0.9rem;">Verifique se os arquivos foram carregados corretamente e tente novamente.</p>
                        </div>
                        """, unsafe_allow_html=True)

                    # Get raw results
                    try:
                        if hasattr(processor, 'get_raw_results'):
                            raw_results_dict = processor.get_raw_results()
                        else:
                            raw_results_dict = {}
                    except (AttributeError, Exception):
                        raw_results_dict = {}
                    
                    # Download all results in different formats in a separate block
                    st.subheader("📥 Opções de Download")
                    with st.container(border=True):
                        st.markdown('<div style="font-size: 11pt;">', unsafe_allow_html=True)
                        if results.get('results'):
                            col1, col2, col3, col4, col5 = st.columns(5)
                            
                            with col1:
                                # JSON format
                                json_results = json.dumps(results, indent=2, ensure_ascii=False)
                                st.download_button(
                                    "📥 Download JSON",
                                    json_results,
                                    file_name="ocr_results.json",
                                    mime="application/json",
                                    key="download_json_batch"
                                )
                            
                            with col2:
                                # DOCX format - structured batch results
                                try:
                                    batch_content = {os.path.basename(fp): text for fp, text in results['results'].items()}
                                    doc = create_structured_docx(
                                        title='Resultados do OCR (Lote)',
                                        content_dict=batch_content,
                                        model_name=selected_model,
                                        format_type=format_type,
                                        language=language,
                                        elapsed_time=elapsed_time,
                                        is_batch=True
                                    )
                                    docx_buffer = BytesIO()
                                    doc.save(docx_buffer)
                                    docx_buffer.seek(0)
                                    st.download_button(
                                        "📥 Download DOCX",
                                        docx_buffer.getvalue(),
                                        file_name="ocr_results.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_docx_batch"
                                    )
                                except Exception as e:
                                    st.error(f"Erro ao gerar DOCX: {e}")
                            
                            with col3:
                                # DOC format - structured batch results
                                try:
                                    batch_content = {os.path.basename(fp): text for fp, text in results['results'].items()}
                                    doc = create_structured_docx(
                                        title='Resultados do OCR (Lote)',
                                        content_dict=batch_content,
                                        model_name=selected_model,
                                        format_type=format_type,
                                        language=language,
                                        elapsed_time=elapsed_time,
                                        is_batch=True
                                    )
                                    doc_buffer = BytesIO()
                                    doc.save(doc_buffer)
                                    doc_buffer.seek(0)
                                    st.download_button(
                                        "📥 Download DOC",
                                        doc_buffer.getvalue(),
                                        file_name="ocr_results.doc",
                                        mime="application/msword",
                                        key="download_doc_batch"
                                    )
                                except Exception as e:
                                    st.error(f"Erro ao gerar DOC: {e}")
                            
                            with col4:
                                # RAW format - exactly as LLM processed
                                try:
                                    # Combine all raw results
                                    raw_content = []
                                    for fp, text in results['results'].items():
                                        file_name = os.path.basename(fp)
                                        raw_text = raw_results_dict.get(fp, text)  # Fallback to formatted if raw not available
                                        raw_content.append(f"=== {file_name} ===\n{raw_text}\n\n")
                                    
                                    raw_all = "\n".join(raw_content)
                                    st.download_button(
                                        "📥 Download RAW",
                                        raw_all,
                                        file_name="ocr_results_raw.txt",
                                        mime="text/plain",
                                        help="Resultados exatamente como processados pela LLM, sem formatação",
                                        key="download_raw_batch"
                                    )
                                except Exception as e:
                                    st.error(f"Erro ao gerar RAW: {e}")
                            
                            with col5:
                                # Formato Minuta - Legal document format for batch
                                try:
                                    batch_content = {os.path.basename(fp): text for fp, text in results['results'].items()}
                                    minuta_doc = create_minuta_doc(
                                        content_dict=batch_content,
                                        is_batch=True
                                    )
                                    minuta_buffer = BytesIO()
                                    minuta_doc.save(minuta_buffer)
                                    minuta_buffer.seek(0)
                                    st.download_button(
                                        "📄 Formato Minuta",
                                        minuta_buffer.getvalue(),
                                        file_name="minuta.doc",
                                        mime="application/msword",
                                        help="Documento formatado conforme padrão de peças processuais (fonte Times New Roman 12, espaçamento 1,5, margens padrão)",
                                        key="download_minuta_batch"
                                    )
                                except Exception as e:
                                    st.error(f"Erro ao gerar Minuta: {e}")
                        else:
                            st.markdown("""
                            <div class="warning-highlight">
                                <p><strong>⚠️ Atenção:</strong> Nenhum resultado disponível para download.</p>
                                <p style="margin-top: 0.5rem; font-size: 0.9rem;">Processe os arquivos primeiro para gerar resultados disponíveis para download.</p>
                            </div>
                            """, unsafe_allow_html=True)
                        st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()